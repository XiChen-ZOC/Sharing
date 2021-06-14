import gevent
from gevent import monkey
monkey.patch_all()
import requests
from lxml import html
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook

# 多进程 多线程 协程

def get_article(keyword):
    url = f'https://pubmed.ncbi.nlm.nih.gov/?term={keyword}&size=100'
    html_data = requests.get(url).text
    selector = html.fromstring(html_data)
    articles = selector.xpath('//*[@id="search-results"]/section/div[1]/div[1]/article')
    tasks_list = []
    for article in articles:
        # get_info(article)
        task = gevent.spawn(get_info, article)
        tasks_list.append(task)
    gevent.joinall(tasks_list)
    return article_lst

def get_info(article):
    # 标题
    title_res = article.xpath('div[2]/div[1]/a//text()')
    # results_lst = []
    # for i in results:
    #     if i.strip() != '':
    #         results_lst.append(i.strip)
    results = [i.strip() for i in title_res if i.strip()]
    title = ' '.join(results).strip('.').strip()
    print(title)

    # 作者
    authors = article.xpath('div[2]/div[1]/div[1]/span[1]/text()')[0]
    # print(authors)

    # 其他信息
    info = article.xpath('div[2]/div[1]/div[1]/span[3]/text()')[0]
    # print(info)
    info_lst = info.split('. ')
    journal, pub_data, doi = info_lst[0], info_lst[1], info_lst[2]
    print(journal)
    print(pub_data)
    print(doi)

    # 详情页
    href = article.xpath('div[2]/div[1]/a//@href')
    url_detail = 'https://pubmed.ncbi.nlm.nih.gov' + href[0]
    html_data_detail = requests.get(url_detail).text
    selector_detail = html.fromstring(html_data_detail)
    res = selector_detail.xpath("//*[@id='enc-abstract']/p//text()")
    results_abstract = [i.strip() for i in res if i.strip()]
    abstract = ' '.join(results_abstract)
    print(abstract)
    print('-' * 20)

    save_word(title, abstract)
    article_lst.append([title, authors, journal, pub_data, doi, abstract])
    return article_lst

def save_word(title, abstract):
    desktop = r'C:\Users\chenx\Desktop'
    wordfile = Document()
    wordfile.add_heading('Abstracts')
    p = wordfile.add_paragraph(abstract)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    wordfile.save(desktop + r'\Data' + r'\{}.docx'.format(
        title.replace('/', '').replace('\\', '').replace(':', '').\
            replace('?', '').replace('<', '').replace('>', '').\
            replace('|', '').replace('*', '').replace('"', '')))

def save_excel(results, keyword):
    desktop = r'C:\Users\chenx\Desktop'
    workbook = Workbook()
    sheet = workbook.active
    header = ['序号','题目','作者','期刊','出版日期','doi号', '摘要']
    sheet.append(header)
    n = 0
    for i in results:
        n += 1
        i.insert(0, n)
        sheet.append(i)
    workbook.save(desktop + r'\Data\检索结果 {}.xlsx'.format(keyword.capitalize()))


if __name__ == '__main__':
    keyword = 'cataract machine learning'

    article_lst = []
    results = get_article(keyword)
    save_excel(results, keyword)


